"""
Document Intelligence Service
Advanced OCR, text extraction, and document classification capabilities
"""

import io
import tempfile
import os
from typing import Dict, Any, List, Optional, Tuple
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import PyPDF2
import docx
from docx import Document
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import logging

from src.core.logging import api_logger as logger


class DocumentIntelligenceService:
    """Advanced document intelligence with OCR and classification"""
    
    def __init__(self):
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        
        # Regulatory document patterns for classification
        self.regulatory_patterns = {
            'rbi_circular': [
                r'reserve bank of india',
                r'rbi/\d{4}-\d{2}',
                r'circular.*rbi',
                r'master.*circular',
                r'prudential.*norms',
                r'banking.*regulation'
            ],
            'compliance_guideline': [
                r'compliance.*guideline',
                r'regulatory.*compliance',
                r'aml.*cft',
                r'know.*your.*customer',
                r'kyc.*norms',
                r'suspicious.*transaction'
            ],
            'risk_management': [
                r'risk.*management',
                r'credit.*risk',
                r'operational.*risk',
                r'market.*risk',
                r'basel.*iii',
                r'capital.*adequacy'
            ],
            'policy_document': [
                r'policy.*document',
                r'internal.*policy',
                r'board.*approved',
                r'policy.*framework',
                r'governance.*policy'
            ],
            'financial_regulation': [
                r'financial.*regulation',
                r'sebi.*regulation',
                r'irdai.*regulation',
                r'nbfc.*guidelines',
                r'banking.*act'
            ]
        }
        
        logger.info("Document Intelligence Service initialized")
    
    async def extract_text_from_file(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """Extract text from various file formats with OCR support"""
        try:
            extracted_text = ""
            confidence = 0.0
            method = "unknown"
            
            if content_type == "text/plain":
                extracted_text = file_content.decode('utf-8')
                confidence = 1.0
                method = "direct_text"
                
            elif content_type == "application/pdf":
                extracted_text, confidence, method = await self._extract_from_pdf(file_content)
                
            elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                extracted_text, confidence, method = await self._extract_from_word(file_content)
                
            elif content_type.startswith("image/"):
                extracted_text, confidence, method = await self._extract_from_image(file_content)
                
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            return {
                "extracted_text": extracted_text,
                "confidence": confidence,
                "method": method,
                "text_length": len(extracted_text),
                "word_count": len(extracted_text.split()) if extracted_text else 0,
                "filename": filename,
                "content_type": content_type
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            raise
    
    async def _extract_from_pdf(self, file_content: bytes) -> Tuple[str, float, str]:
        """Extract text from PDF with fallback to OCR"""
        try:
            # First try direct text extraction
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text.strip(), 0.95, "pdf_direct"
            
            # If no text found, try OCR on PDF pages
            logger.info("No text found in PDF, attempting OCR")
            return await self._ocr_pdf_pages(file_content)
            
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {e}")
            # Fallback to OCR
            return await self._ocr_pdf_pages(file_content)
    
    async def _extract_from_word(self, file_content: bytes) -> Tuple[str, float, str]:
        """Extract text from Word documents"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip(), 0.98, "word_direct"
            
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            raise
    
    async def _extract_from_image(self, file_content: bytes) -> Tuple[str, float, str]:
        """Extract text from images using OCR"""
        try:
            # Open and preprocess image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Enhance image for better OCR
            image = self._enhance_image_for_ocr(image)
            
            # Perform OCR
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,!?@#$%^&*()_+-=[]{}|;:\'\"<>/\\ '
            
            text = pytesseract.image_to_string(image, config=custom_config)
            
            # Get confidence data
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return text.strip(), avg_confidence / 100.0, "ocr_image"
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise
    
    async def _ocr_pdf_pages(self, file_content: bytes) -> Tuple[str, float, str]:
        """Perform OCR on PDF pages"""
        try:
            # This would require pdf2image library to convert PDF to images
            # For now, return a placeholder
            logger.warning("PDF OCR not fully implemented - requires pdf2image library")
            return "OCR extraction from PDF pages not yet implemented", 0.5, "pdf_ocr_placeholder"
            
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            raise
    
    def _enhance_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Enhance image quality for better OCR results"""
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # Apply slight blur to reduce noise
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            return image
            
        except Exception as e:
            logger.warning(f"Image enhancement failed: {e}")
            return image
    
    async def classify_document(self, text: str, categories: Optional[List[str]] = None) -> Dict[str, Any]:
        """Classify document based on content using pattern matching and ML"""
        try:
            if not text.strip():
                return {
                    "predicted_category": "unknown",
                    "confidence": 0.0,
                    "method": "no_text",
                    "all_scores": []
                }
            
            text_lower = text.lower()
            
            # Pattern-based classification
            pattern_scores = {}
            for category, patterns in self.regulatory_patterns.items():
                score = 0
                matches = 0
                
                for pattern in patterns:
                    if re.search(pattern, text_lower):
                        matches += 1
                        score += 1
                
                # Normalize score
                if patterns:
                    pattern_scores[category] = (score / len(patterns)) * (matches / len(patterns))
            
            # Find best match
            if pattern_scores:
                best_category = max(pattern_scores, key=pattern_scores.get)
                best_score = pattern_scores[best_category]
                
                # Convert to list format
                all_scores = [
                    {"category": cat, "score": score, "confidence": score}
                    for cat, score in sorted(pattern_scores.items(), key=lambda x: x[1], reverse=True)
                ]
                
                return {
                    "predicted_category": best_category,
                    "confidence": best_score,
                    "method": "pattern_matching",
                    "all_scores": all_scores
                }
            
            # Fallback to generic classification
            return {
                "predicted_category": "general_document",
                "confidence": 0.3,
                "method": "fallback",
                "all_scores": [{"category": "general_document", "score": 0.3, "confidence": 0.3}]
            }
            
        except Exception as e:
            logger.error(f"Document classification failed: {e}")
            raise
    
    async def analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and extract metadata"""
        try:
            lines = text.split('\n')
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            # Extract potential headers (lines with fewer than 100 chars and title case)
            headers = []
            for line in lines[:20]:  # Check first 20 lines
                line = line.strip()
                if line and len(line) < 100 and (line.isupper() or line.istitle()):
                    headers.append(line)
            
            # Extract dates
            date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
            dates = re.findall(date_pattern, text)
            
            # Extract numbers/references
            ref_pattern = r'\b[A-Z]{2,}/\d{4}-\d{2}/\d+\b|\b\d{4}/\d+\b'
            references = re.findall(ref_pattern, text)
            
            return {
                "total_lines": len(lines),
                "total_paragraphs": len(paragraphs),
                "potential_headers": headers[:5],  # Top 5 headers
                "extracted_dates": dates[:10],     # Top 10 dates
                "extracted_references": references[:5],  # Top 5 references
                "avg_paragraph_length": sum(len(p) for p in paragraphs) / len(paragraphs) if paragraphs else 0,
                "structure_score": min(1.0, len(headers) * 0.2 + len(paragraphs) * 0.1)
            }
            
        except Exception as e:
            logger.error(f"Document structure analysis failed: {e}")
            return {}
